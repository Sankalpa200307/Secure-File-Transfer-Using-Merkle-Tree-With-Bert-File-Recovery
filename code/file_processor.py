'''
Enhanced File Processor for Multiple Formats
Supports text extraction from PDF, DOCX, and other text formats
'''

import os
import logging
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class FileProcessor:
    """
    Enhanced file processor for multiple document formats
    """
    
    def __init__(self):
        self.supported_extensions = {
            # Text formats
            '.txt': self._read_text_file,
            '.md': self._read_text_file,
            '.py': self._read_text_file,
            '.js': self._read_text_file,
            '.html': self._read_text_file,
            '.css': self._read_text_file,
            '.json': self._read_text_file,
            '.xml': self._read_text_file,
            '.csv': self._read_text_file,
            # Document formats
            '.pdf': self._read_pdf_file,
            '.docx': self._read_docx_file,
            '.doc': self._read_docx_file,  # Will try to convert
        }
    
    def is_supported_format(self, file_path):
        """
        Check if file format is supported for text extraction
        """
        _, ext = os.path.splitext(file_path.lower())
        return ext in self.supported_extensions
    
    def get_supported_formats(self):
        """
        Get list of supported file formats
        """
        return list(self.supported_extensions.keys())
    
    def extract_text(self, file_path):
        """
        Extract text from file regardless of format
        
        Args:
            file_path: Path to the file
            
        Returns:
            str: Extracted text content, or None if failed
        """
        try:
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return None
            
            _, ext = os.path.splitext(file_path.lower())
            
            if ext not in self.supported_extensions:
                logger.warning(f"Unsupported file format: {ext}")
                return None
            
            # Extract text using appropriate method
            text = self.supported_extensions[ext](file_path)
            
            if text:
                logger.info(f"✅ Successfully extracted text from {file_path}")
                logger.info(f"   Text length: {len(text)} characters")
                return text
            else:
                logger.error(f"❌ Failed to extract text from {file_path}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {e}")
            return None
    
    def _read_text_file(self, file_path):
        """
        Read text from plain text files
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
                
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return None
    
    def _read_pdf_file(self, file_path):
        """
        Extract text from PDF files using multiple methods
        """
        try:
            # Method 1: Try pdfplumber (better for complex PDFs)
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    
                    if text.strip():
                        logger.info(f"✅ Extracted text using pdfplumber: {len(text)} characters")
                        return text.strip()
            except ImportError:
                logger.warning("pdfplumber not available, trying PyPDF2")
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
            
            # Method 2: Try PyPDF2 (fallback)
            try:
                import PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                    
                    if text.strip():
                        logger.info(f"✅ Extracted text using PyPDF2: {len(text)} characters")
                        return text.strip()
            except ImportError:
                logger.error("PyPDF2 not available for PDF processing")
            except Exception as e:
                logger.error(f"PyPDF2 failed: {e}")
            
            logger.error(f"❌ Failed to extract text from PDF: {file_path}")
            return None
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            return None
    
    def _read_docx_file(self, file_path):
        """
        Extract text from DOCX files
        """
        try:
            import docx
            
            doc = docx.Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            if text.strip():
                logger.info(f"✅ Extracted text from DOCX: {len(text)} characters")
                return text.strip()
            else:
                logger.warning(f"❌ No text found in DOCX file: {file_path}")
                return None
                
        except ImportError:
            logger.error("python-docx not available for DOCX processing")
            return None
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            return None
    
    def save_text_to_file(self, text, output_path):
        """
        Save extracted text to a file
        
        Args:
            text: Text content to save
            output_path: Path where to save the text file
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            logger.info(f"✅ Text saved to: {output_path}")
            return True
        except Exception as e:
            logger.error(f"Error saving text to {output_path}: {e}")
            return False
    
    def get_file_info(self, file_path):
        """
        Get information about a file
        
        Args:
            file_path: Path to the file
            
        Returns:
            dict: File information
        """
        try:
            if not os.path.exists(file_path):
                return None
            
            stat = os.stat(file_path)
            _, ext = os.path.splitext(file_path.lower())
            
            info = {
                'path': file_path,
                'name': os.path.basename(file_path),
                'extension': ext,
                'size': stat.st_size,
                'is_supported': self.is_supported_format(file_path),
                'supported_formats': self.get_supported_formats()
            }
            
            return info
            
        except Exception as e:
            logger.error(f"Error getting file info for {file_path}: {e}")
            return None

def main():
    """
    Test the file processor
    """
    print("🧪 Testing Enhanced File Processor")
    print("=" * 40)
    
    processor = FileProcessor()
    
    print(f"Supported formats: {', '.join(processor.get_supported_formats())}")
    
    # Test with a sample text file
    test_content = """
    This is a test document for the enhanced file processor.
    It should be able to handle multiple file formats including PDF and DOCX.
    """
    
    # Create test files
    test_files = [
        ("test.txt", test_content),
        ("test.md", f"# Test Document\n\n{test_content}"),
        ("test.json", '{"title": "Test Document", "content": "This is test content"}')
    ]
    
    for filename, content in test_files:
        print(f"\n📝 Testing: {filename}")
        
        # Create file
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Test extraction
        extracted_text = processor.extract_text(filename)
        
        if extracted_text:
            print(f"✅ Successfully extracted {len(extracted_text)} characters")
        else:
            print(f"❌ Failed to extract text")
        
        # Clean up
        if os.path.exists(filename):
            os.remove(filename)
    
    print("\n🎉 File processor test completed!")

if __name__ == "__main__":
    main() 